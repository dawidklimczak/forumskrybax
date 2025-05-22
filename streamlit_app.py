import streamlit as st
import os
import assemblyai as aai
from docx import Document
import tempfile
import io

# Konfiguracja strony Streamlit
st.set_page_config(page_title="Transkrypcja Audio", page_icon="🎤")
st.title("Transkrypcja plików audio")

# Inicjalizacja stanu sesji dla transkrypcji
if 'transcription_done' not in st.session_state:
    st.session_state.transcription_done = False
if 'full_transcript' not in st.session_state:
    st.session_state.full_transcript = ""
if 'utterances_data' not in st.session_state:
    st.session_state.utterances_data = []

# Inicjalizacja AssemblyAI
def initialize_assemblyai():
    try:
        api_key = st.secrets.get("ASSEMBLYAI_API_KEY")
        if not api_key:
            api_key = os.environ.get("ASSEMBLYAI_API_KEY")
        
        if not api_key:
            raise ValueError("Nie znaleziono klucza API AssemblyAI")
            
        aai.settings.api_key = api_key
        return True
    except Exception as e:
        st.error(f"Błąd podczas inicjalizacji API AssemblyAI: {str(e)}")
        return False

# Inicjalizacja klienta AssemblyAI
if not initialize_assemblyai():
    st.stop()

def transcribe_audio(file_path, speakers_expected=2):
    """
    Transkrybuje plik audio używając AssemblyAI.
    
    Args:
        file_path (str): Ścieżka do pliku audio
        speakers_expected (int): Oczekiwana liczba rozmówców
        
    Returns:
        tuple: (lista wypowiedzi, pełna transkrypcja jako tekst)
    """
    try:
        # Konfiguracja transkrypcji
        config = aai.TranscriptionConfig(
            language_code="pl",  # Język polski
            speaker_labels=True,  # Włączenie etykiet mówców
            speakers_expected=speakers_expected   # Oczekiwana liczba rozmówców
        )
        
        # Transkrypcja
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(file_path, config=config)
        
        # Sprawdzenie czy transkrypcja się powiodła
        if transcript.status == aai.TranscriptStatus.error:
            raise Exception(f"Błąd transkrypcji: {transcript.error}")
        
        # Przetwarzanie transkrypcji
        utterances_data = []
        full_text_parts = []
        
        for utterance in transcript.utterances:
            utterance_info = {
                "speaker": f"Rozmówca {utterance.speaker}",
                "text": utterance.text,
                "start": utterance.start,
                "end": utterance.end
            }
            utterances_data.append(utterance_info)
            full_text_parts.append(f"{utterance_info['speaker']}: {utterance_info['text']}")
        
        full_transcript = "\n\n".join(full_text_parts)
        
        return utterances_data, full_transcript
    
    except Exception as e:
        st.error(f"Błąd podczas transkrypcji: {str(e)}")
        raise

def save_to_word(utterances_data, filename="transkrypcja.docx"):
    """Zapisuje transkrypcję z podziałem na mówców do pliku Word"""
    doc = Document()
    doc.add_heading('Transkrypcja nagrania', 0)
    
    for utterance in utterances_data:
        # Dodaj nagłówek z mówcą
        speaker_paragraph = doc.add_paragraph()
        speaker_run = speaker_paragraph.add_run(f"{utterance['speaker']}:")
        speaker_run.bold = True
        
        # Dodaj tekst wypowiedzi
        doc.add_paragraph(utterance['text'])
        
        # Dodaj pustą linię dla czytelności
        doc.add_paragraph()
    
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

# Interface użytkownika
st.sidebar.header("Ustawienia transkrypcji")
speakers_expected = st.sidebar.number_input(
    "Oczekiwana liczba rozmówców",
    min_value=1,
    max_value=10,
    value=2,
    help="Podaj przewidywaną liczbę osób mówiących w nagraniu"
)

uploaded_file = st.file_uploader("Wybierz plik audio", type=['mp3', 'wav', 'm4a', 'ogg', 'flac'])

if uploaded_file:
    st.info(f"Wybrano plik: {uploaded_file.name}")
    st.info(f"Rozmiar pliku: {uploaded_file.size / (1024*1024):.2f} MB")
    
    # Przycisk do rozpoczęcia transkrypcji
    if st.button("Rozpocznij transkrypcję"):
        with st.spinner('Trwa transkrypcja nagrania (to może zająć kilka minut)...'):
            # Zapisz uploadowany plik tymczasowo
            temp_input = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{uploaded_file.name.split(".")[-1]}')
            temp_input.write(uploaded_file.getvalue())
            temp_input.close()
            
            try:
                # Transkrypcja
                utterances_data, full_transcript = transcribe_audio(temp_input.name, speakers_expected)
                
                # Zapisz wyniki do session_state
                st.session_state.utterances_data = utterances_data
                st.session_state.full_transcript = full_transcript
                st.session_state.transcription_done = True
                
                st.success("Transkrypcja zakończona pomyślnie!")
                
            except Exception as e:
                st.error(f"Wystąpił błąd podczas transkrypcji: {str(e)}")
            
            finally:
                # Usuń tymczasowy plik
                if os.path.exists(temp_input.name):
                    os.unlink(temp_input.name)

    # Wyświetl transkrypcję i przycisk do pobrania tylko jeśli transkrypcja została wykonana
    if st.session_state.transcription_done and st.session_state.full_transcript:
        st.subheader("Transkrypcja:")
        
        # Opcja wyświetlania
        display_mode = st.radio(
            "Sposób wyświetlania:",
            ["Pełna transkrypcja", "Podział na wypowiedzi"],
            horizontal=True
        )
        
        if display_mode == "Pełna transkrypcja":
            st.text_area("", st.session_state.full_transcript, height=400)
        else:
            # Wyświetl z podziałem na wypowiedzi
            for i, utterance in enumerate(st.session_state.utterances_data):
                with st.expander(f"{utterance['speaker']} ({utterance['start']/1000:.1f}s - {utterance['end']/1000:.1f}s)"):
                    st.write(utterance['text'])
        
        # Przycisk do pobrania pliku Word
        doc_buffer = save_to_word(st.session_state.utterances_data)
        st.download_button(
            label="Pobierz jako dokument Word",
            data=doc_buffer,
            file_name="transkrypcja.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Statystyki
        st.subheader("Statystyki:")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Liczba wypowiedzi", len(st.session_state.utterances_data))
        
        with col2:
            unique_speakers = len(set([u['speaker'] for u in st.session_state.utterances_data]))
            st.metric("Liczba rozmówców", unique_speakers)
        
        with col3:
            if st.session_state.utterances_data:
                total_duration = st.session_state.utterances_data[-1]['end'] / 1000
                st.metric("Czas trwania", f"{total_duration:.1f}s")